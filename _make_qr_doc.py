"""Erzeugt Mathepfade_QR.docx mit QR-Code fuer Mathepfade."""
from io import BytesIO
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

URL = "https://nachbar-blip.github.io/Mathepfade/"


def make_qr_png(url: str) -> BytesIO:
    qr = qrcode.QRCode(version=None, box_size=12, border=2,
                       error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    title = doc.add_heading("Mathepfade", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Adaptives Mathe-Training · Thueringen · Klasse 7-12")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    doc.add_paragraph()

    p_qr = doc.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr.add_run().add_picture(make_qr_png(URL), width=Cm(9.0))

    p_url = doc.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_url = p_url.add_run(URL)
    r_url.font.size = Pt(11)
    r_url.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        "90 Trainer, je 36 Aufgaben in 6 Schwierigkeitsstufen.\n"
        "Laeuft direkt im Browser — keine Anmeldung, keine App noetig."
    )
    ir.font.size = Pt(11)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Scanne mit dem Handy und lege los.")
    fr.italic = True
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = Path(__file__).parent / "Mathepfade_QR.docx"
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
